"""
Ingest downloaded DLGF documents and write a text extraction parquet.

Reads docs/ layout and metadata.json produced by get_docs.py.
Outputs documents.parquet with one row per document.

Columns: doc_id, source, doc_type, text, char_count, retrieved_at
"""

import hashlib
import json
import warnings
from datetime import datetime, timezone
from pathlib import Path

import openpyxl
import pandas as pd
import pdfplumber
from docx import Document

warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")

METADATA_FILE = Path("metadata.json")
DOCS_DIR = Path("docs")
OUTPUT_FILE = Path("documents.parquet")


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------

def extract_pdf(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text.strip())
    return "\n\n".join(parts)


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def extract_xlsx(path: Path) -> str:
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets: list[str] = []
    for ws in wb.worksheets:
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                rows.append("\t".join(cells))
        if rows:
            sheets.append(f"[{ws.title}]\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets)


EXTRACTORS = {
    "pdf": extract_pdf,
    "docx": extract_docx,
    "doc": extract_docx,
    "xlsx": extract_xlsx,
    "xls": extract_xlsx,
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def make_doc_id(url: str) -> str:
    return hashlib.md5(url.encode()).hexdigest()


def resolve_doc_type(meta: dict) -> str:
    """Use semantic doc_type from filename parse; fall back to file extension."""
    return meta.get("doc_type") or meta["file_type"].upper()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    all_docs = json.loads(METADATA_FILE.read_text(encoding="utf-8"))
    downloaded = [d for d in all_docs if d.get("downloaded")]
    print(f"Documents to ingest: {len(downloaded)}")

    retrieved_at = datetime.now(timezone.utc).isoformat()

    records: list[dict] = []
    errors = 0

    for i, doc in enumerate(downloaded, 1):
        path = DOCS_DIR / doc["local_path"]
        file_type = doc["file_type"]
        extractor = EXTRACTORS.get(file_type)

        label = doc["filename"][:60]
        print(f"  [{i:>4}/{len(downloaded)}] ", end="", flush=True)

        if extractor is None:
            print(f"SKIP (unsupported: {file_type})  {label}")
            continue

        if not path.exists():
            print(f"MISS (file not found)  {label}")
            errors += 1
            continue

        try:
            text = extractor(path)
            char_count = len(text)
            print(f"{'EMPTY' if char_count == 0 else 'OK   '} ({char_count:>7,} chars)  {label}")
        except Exception as exc:
            print(f"ERR   {exc}  {label}")
            text = ""
            char_count = 0
            errors += 1

        records.append(
            {
                "doc_id": make_doc_id(doc["url"]),
                "source": doc["url"],
                "doc_type": resolve_doc_type(doc),
                "text": text,
                "char_count": char_count,
                "retrieved_at": retrieved_at,
            }
        )

    df = pd.DataFrame(
        records,
        columns=["doc_id", "source", "doc_type", "text", "char_count", "retrieved_at"],
    )
    df.to_parquet(OUTPUT_FILE, index=False)

    print(f"\nWrote {len(df):,} rows -> {OUTPUT_FILE}")
    print(f"  Errors / missing: {errors}")

    empty = (df["char_count"] == 0).sum()
    if empty:
        print(
            f"  Warning: {empty} documents have no extracted text "
            "(likely scanned/image PDFs — consider adding OCR in a later step)"
        )

    print("\nBreakdown by doc_type:")
    for doc_type, count in df["doc_type"].value_counts().items():
        print(f"  {doc_type:<15} {count:>4}")


if __name__ == "__main__":
    main()